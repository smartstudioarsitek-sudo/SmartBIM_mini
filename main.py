import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from io import BytesIO
import json
import re
import time
from PIL import Image
import docx  # library: python-docx
import zipfile
from pptx import Presentation # library: python-pptx
from streamlit_drawable_canvas import st_canvas

# --- AI & GOOGLE LIBRARIES ---
import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold

# --- IMPORT MODULE ENGINEERING LOKAL ---
# Pastikan file-file ini ada di folder yang sama
import libs_sni as sni
import libs_ahsp as ahsp
import libs_bim_importer as bim
import libs_pondasi as fdn
import libs_geoteknik as geo
import libs_export as exp
import libs_baja as steel
import libs_gempa as quake

# --- IMPORT BACKEND DATABASE ---
try:
    from backend_enginex import EnginexBackend
except ImportError:
    st.error("⚠️ File 'backend_enginex.py' tidak ditemukan. Upload file tersebut agar fitur Chat berfungsi.")
    st.stop()

# ==========================================
# 1. KONFIGURASI HALAMAN & CSS
# ==========================================
st.set_page_config(
    page_title="IndoBIM x ENGINEX Ultimate", 
    layout="wide", 
    page_icon="🏗️",
    initial_sidebar_state="expanded"
)

# --- CSS CUSTOM UNTUK TAMPILAN PROFESIONAL ---
st.markdown("""
<style>
    .main-header {font-size: 28px; font-weight: bold; color: #1E3A8A; margin-bottom: 5px;}
    .sub-header {font-size: 18px; font-weight: bold; color: #444;}
    
    /* Tabs Customization */
    .stTabs [data-baseweb="tab-list"] { gap: 2px; }
    .stTabs [data-baseweb="tab"] {
        height: 50px; white-space: pre-wrap; background-color: #f0f2f6; border-radius: 4px 4px 0 0; gap: 1px; padding-top: 10px; padding-bottom: 10px;
    }
    .stTabs [aria-selected="true"] { background-color: #1E3A8A; color: white; }
    
    /* Chat Styling */
    .stChatMessage .avatar {background-color: #1E3A8A; color: white;}
    .auto-pilot-msg { background-color: #e0f7fa; border-left: 5px solid #00acc1; padding: 10px; margin-bottom: 10px; border-radius: 5px; color: #006064; font-weight: bold; }
    
    /* Button Styling */
    div.stButton > button {width: 100%; border-radius: 6px; font-weight: 600;}
</style>
""", unsafe_allow_html=True)

# ==========================================
# 2. INISIALISASI STATE (MEMORI APLIKASI)
# ==========================================

# A. State Kalkulator (IndoBIM) - Agar data tidak hilang saat interaksi
if 'geo' not in st.session_state: st.session_state['geo'] = {'L': 6.0, 'b': 250, 'h': 400}
if 'structure' not in st.session_state: st.session_state['structure'] = {}
if 'pondasi' not in st.session_state: st.session_state['pondasi'] = {}
if 'geotech' not in st.session_state: st.session_state['geotech'] = {}
if 'arsitek_mep' not in st.session_state: st.session_state['arsitek_mep'] = {} # Tempat data IFC
if 'bim_loads' not in st.session_state: st.session_state['bim_loads'] = 0
if 'drawing' not in st.session_state: st.session_state['drawing'] = {}

# State Report (Untuk PDF/Laporan)
for k in ['report_struk', 'report_baja', 'report_gempa', 'report_geo']:
    if k not in st.session_state: st.session_state[k] = {}

# B. State AI Chat (ENGINEX)
if 'backend' not in st.session_state: st.session_state.backend = EnginexBackend()
db = st.session_state.backend

if 'processed_files' not in st.session_state: st.session_state.processed_files = set()
if 'current_expert_active' not in st.session_state: st.session_state.current_expert_active = "👑 The GEMS Grandmaster"

# ==========================================
# 3. HELPER FUNCTIONS (LOGIKA PENDUKUNG)
# ==========================================

@st.cache_resource
def get_available_models_from_google(api_key_trigger):
    """Mendapatkan daftar model yang TERSEDIA untuk API Key user (Anti Error 404)"""
    try:
        model_list = []
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                model_list.append(m.name)
        # Urutkan agar Flash/Pro ada di atas
        model_list.sort(key=lambda x: 'gemini' not in x) 
        return model_list, None
    except Exception as e:
        return [], str(e)

def get_project_summary_context():
    """
    Mengambil rangkuman data teknis dari Tab Kalkulator & BIM 
    untuk dikirim sebagai 'Context' ke AI.
    """
    summary = "DATA TEKNIS PROYEK SAAT INI (Dari Kalkulator IndoBIM):\n"
    
    # 1. STRUKTUR BETON
    if st.session_state.get('report_struk'):
        s = st.session_state['report_struk']
        summary += f"- Beton: Dimensi {s.get('Dimensi')}, Mu={s.get('Mu')} kNm, Perlu Tulangan={s.get('Tulangan')}\n"
    
    # 2. STRUKTUR BAJA
    if st.session_state.get('report_baja'):
        b = st.session_state['report_baja']
        summary += f"- Baja: Profil {b.get('Profil')}, Ratio={b.get('Ratio')}, Status={b.get('Status')}\n"
    
    # 3. GEOTEKNIK
    if st.session_state.get('report_geo'):
        g = st.session_state['report_geo']
        summary += f"- Geoteknik: SF Talud={g.get('Talud_SF')}, Qall Pile={g.get('Pile_Qall')} kN\n"
    
    # 4. GAMBAR CANVAS (MANUAL)
    if st.session_state.get('drawing'):
        d = st.session_state['drawing']
        summary += f"- Estimasi Gambar Manual: Dinding {d.get('vol_dinding')} m2, Beton {d.get('vol_beton')} m3\n"
    
    # 5. DATA BIM/IFC (PENTING AGAR AI TAHU DATA IFC)
    if st.session_state.get('arsitek_mep'):
        m = st.session_state['arsitek_mep']
        summary += f"\n[DATA HASIL SCAN BIM/IFC 3D]\n"
        summary += f"- Luas Dinding Total: {m.get('Luas Dinding (m2)', 0)} m2\n"
        summary += f"- Jumlah Pintu: {m.get('Jumlah Pintu (Unit)', 0)} unit\n"
        summary += f"- Jumlah Jendela: {m.get('Jumlah Jendela (Unit)', 0)} unit\n"
        summary += f"- Panjang Pipa MEP: {m.get('Panjang Pipa/Duct (m\')', 0)} m\n"
    
    return summary

def create_docx_from_text(text_content):
    """Membuat file Word dari jawaban AI"""
    try:
        doc = docx.Document()
        doc.add_heading('Laporan Konsultasi AI - IndoBIM', 0)
        for line in text_content.split('\n'):
            clean = line.strip()
            if clean.startswith('## '): doc.add_heading(clean.replace('## ', ''), level=2)
            elif clean.startswith('### '): doc.add_heading(clean.replace('### ', ''), level=3)
            elif clean.startswith('- ') or clean.startswith('* '): 
                try: doc.add_paragraph(clean, style='List Bullet')
                except: doc.add_paragraph(clean)
            elif clean: doc.add_paragraph(clean)
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio
    except: return None

def extract_table_to_excel(text_content):
    """Mendeteksi tabel Markdown dan convert ke Excel"""
    try:
        lines = text_content.split('\n')
        table_data = []
        for line in lines:
            stripped = line.strip()
            if "|" in stripped:
                if set(stripped.replace('|', '').replace('-', '').replace(' ', '')) == set(): continue
                row_cells = [c.strip() for c in stripped.split('|')]
                if stripped.startswith('|'): row_cells = row_cells[1:]
                if stripped.endswith('|'): row_cells = row_cells[:-1]
                if row_cells: table_data.append(row_cells)
        
        if len(table_data) < 2: return None
        headers = table_data[0]
        data_rows = table_data[1:]
        df = pd.DataFrame(data_rows, columns=headers)
        
        output = BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name='Data_AI')
        output.seek(0)
        return output
    except: return None

def execute_generated_code(code_str):
    """Menjalankan kode Python yang digenerate AI (misal: Plotting)"""
    try:
        local_vars = {"pd": pd, "np": np, "plt": plt, "st": st}
        exec(code_str, {}, local_vars)
        return True
    except Exception as e:
        st.error(f"⚠️ Gagal Render Grafik: {e}")
        return False

# ==========================================
# 4. DEFINISI PERSONA (AI EXPERTS)
# ==========================================
PLOT_INSTRUCTION = """
[ATURAN VISUALISASI]: Jika diminta grafik, tulis kode Python (matplotlib) dalam blok ```python. 
Akhiri kode dengan `st.pyplot(plt.gcf())`. Jangan pakai `plt.show()`.
"""

gems_persona = {
    "👑 The GEMS Grandmaster": f"""ANDA ADALAH "THE GEMS GRANDMASTER" (Project Director). 
    Menjawab dengan wawasan Multidisiplin (Teknis, Hukum, Biaya, Agama). {PLOT_INSTRUCTION}""",
    
    "👔 Project Manager (PM)": "ANDA ADALAH SENIOR PROJECT DIRECTOR. Fokus: Strategi, Risiko, Stakeholder.",
    
    "📝 Drafter Laporan DED": "ANDA ADALAH LEAD TECHNICAL WRITER. Fokus: Struktur Laporan PUPR (Pendahuluan, Antara, Akhir).",
    
    "⚖️ Ahli Legal & Kontrak": "ANDA ADALAH AHLI HUKUM KONSTRUKSI (FIDIC/LPJK). Fokus: Klaim, Sengketa, Adendum.",
    
    "🕌 Dewan Syariah": "ANDA ADALAH ULAMA FIQIH BANGUNAN. Fokus: Arah Kiblat, Akad Istishna, Adab Tetangga.",
    
    "🌾 Ahli IKSI-PAI": f"ANDA ADALAH AHLI IRIGASI (IKSI/PAI). Fokus: Kinerja Sistem Irigasi. {PLOT_INSTRUCTION}",
    
    "🌊 Ahli Bangunan Air": f"ANDA ADALAH HYDRAULIC ENGINEER. Fokus: Bendung, Saluran, Pintu Air. {PLOT_INSTRUCTION}",
    
    "🌧️ Ahli Hidrologi": f"ANDA ADALAH HYDROLOGIST. Fokus: Curah Hujan, Banjir Rencana. {PLOT_INSTRUCTION}",
    
    "🏖️ Ahli Teknik Pantai": f"ANDA ADALAH COASTAL ENGINEER. Fokus: Breakwater, Pasang Surut. {PLOT_INSTRUCTION}",
    
    "🏗️ Ahli Struktur (Gedung)": f"ANDA ADALAH STRUCTURAL ENGINEER. Fokus: Beton, Baja, Gempa SNI 1726. {PLOT_INSTRUCTION}",
    
    "🪨 Ahli Geoteknik": f"ANDA ADALAH GEOTECHNICAL ENGINEER. Fokus: Daya Dukung Tanah, Sondir, Longsor. {PLOT_INSTRUCTION}",
    
    "🛣️ Ahli Jalan & Jembatan": f"ANDA ADALAH HIGHWAY ENGINEER. Fokus: Geometrik Jalan, Perkerasan. {PLOT_INSTRUCTION}",
    
    "🌍 Ahli Geodesi & GIS": "ANDA ADALAH GEOMATICS ENGINEER. Fokus: Topografi, Kontur, Peta.",
    
    "🏛️ Senior Architect": "ANDA ADALAH ARSITEK UTAMA. Fokus: Desain, Estetika, Fungsi Ruang.",
    
    "🌳 Landscape Architect": "ANDA ADALAH ARSITEK LANSKAP. Fokus: Taman, RTH, Hardscape.",
    
    "🎨 The Visionary Architect": "ANDA ADALAH VISUALIZER. Fokus: Membuat Prompt Gambar untuk AI Image Generator.",
    
    "🌍 Ahli Planologi": "ANDA ADALAH URBAN PLANNER. Fokus: Tata Ruang, Zonasi Wilayah.",
    
    "🏭 Ahli Proses Industri": "ANDA ADALAH PROCESS ENGINEER. Fokus: PFD, P&ID, Pabrik.",
    
    "📜 Ahli AMDAL": "ANDA ADALAH KETUA TIM AMDAL. Fokus: UKL-UPL, Dampak Lingkungan.",
    
    "♻️ Ahli Teknik Lingkungan": "ANDA ADALAH SANITARY ENGINEER. Fokus: IPAL, Sampah, Air Bersih.",
    
    "⛑️ Ahli K3 Konstruksi": "ANDA ADALAH SAFETY MANAGER. Fokus: CSMS, APD, Risiko Kecelakaan.",
    
    "💻 Lead Engineering Developer": f"ANDA ADALAH FULL-STACK DEV. Fokus: Python, Streamlit, Coding. {PLOT_INSTRUCTION}",
    
    "📐 CAD & BIM Automator": "ANDA ADALAH BIM MANAGER. Fokus: Revit, IFC, Automation.",
    
    "🖥️ Instruktur Software": "ANDA ADALAH TRAINER. Fokus: Tutorial SAP2000, HEC-RAS, Civil 3D.",
    
    "💰 Ahli Estimator (RAB)": "ANDA ADALAH QUANTITY SURVEYOR. Fokus: AHSP, Volume, Biaya Proyek.",
    
    "💵 Ahli Keuangan Proyek": f"ANDA ADALAH FINANCE MANAGER. Fokus: NPV, IRR, ROI. {PLOT_INSTRUCTION}",
    
    "📜 Ahli Perizinan": "ANDA ADALAH KONSULTAN PERIZINAN. Fokus: PBG, SLF, KRK.",
    
    "🤖 The Enginex Architect": "ANDA ADALAH SYSTEM ADMIN APLIKASI INI."
}

def get_auto_pilot_decision(user_query, model_name):
    """Memilih ahli yang tepat secara otomatis menggunakan AI"""
    try:
        router_model = genai.GenerativeModel(model_name)
        list_ahli = list(gems_persona.keys())
        router_prompt = f"""
        Pilih SATU ahli dari daftar berikut untuk menjawab pertanyaan: "{user_query}"
        Daftar: {list_ahli}
        Output: HANYA nama ahli persis. Jika ragu, pilih '👑 The GEMS Grandmaster'.
        """
        response = router_model.generate_content(router_prompt)
        suggested = response.text.strip()
        if suggested in list_ahli: return suggested
        return "👑 The GEMS Grandmaster"
    except:
        return "👑 The GEMS Grandmaster"

# ==========================================
# 5. SIDEBAR UTAMA (API & MODEL)
# ==========================================
with st.sidebar:
    st.title("🏗️ SYSTEM CONTROLLER")
    
    # --- SETUP API KEY ---
    api_key_input = st.text_input("🔑 Google API Key:", type="password")
    if api_key_input:
        raw_key = api_key_input
        st.caption("ℹ️ Menggunakan Key Manual")
    else:
        raw_key = st.secrets.get("GOOGLE_API_KEY")
    
    if not raw_key:
        st.warning("⚠️ Masukkan API Key Google.")
        st.stop()
        
    clean_api_key = raw_key.strip()
    try:
        genai.configure(api_key=clean_api_key)
    except Exception as e:
        st.error(f"Config Error: {e}")

    # --- MODEL SELECTOR (ANTI ERROR 404) ---
    st.markdown("### 🧠 Pilih Otak AI")
    available_models, err = get_available_models_from_google(clean_api_key)
    
    if err:
        st.error(f"Gagal memuat model: {err}")
        selected_model_name = "gemini-1.5-flash" # Fallback
    elif not available_models:
        st.warning("Tidak ada model tersedia untuk Key ini.")
        selected_model_name = "gemini-1.5-flash"
    else:
        # Cari default flash jika ada
        default_idx = 0
        for i, m in enumerate(available_models):
            if "flash" in m: default_idx = i; break
        selected_model_name = st.selectbox("Model:", available_models, index=default_idx)

    st.divider()
    
    # --- PILIH MODE APLIKASI ---
    app_mode = st.radio(
        "Modul Utama:", 
        ["🧮 Kalkulator Teknik (Tools)", "🤖 Konsultan AI (Chat)"], 
        index=0
    )
    
    st.divider()
    
    # --- SETTING PROYEK ---
    st.markdown("### 📁 Database")
    existing_projects = db.daftar_proyek()
    proj_mode = st.radio("Opsi:", ["Baru", "Buka"], horizontal=True, label_visibility="collapsed")
    if proj_mode == "Baru":
        nama_proyek = st.text_input("Nama Proyek:", "Proyek Baru")
    else:
        nama_proyek = st.selectbox("Pilih Proyek:", existing_projects) if existing_projects else "Belum ada"

    st.divider()
    
    # --- PARAMETER TEKNIS ---
    with st.expander("⚙️ Parameter & Harga"):
        fc_in = st.number_input("Mutu Beton f'c (MPa)", 25)
        fy_in = st.number_input("Mutu Besi fy (MPa)", 400)
        gamma_tanah = st.number_input("Berat Tanah (kN/m3)", 18.0)
        phi_tanah = st.number_input("Phi Tanah (deg)", 30.0)
        c_tanah = st.number_input("Kohesi (kN/m2)", 5.0)
        sigma_tanah = st.number_input("Daya Dukung (kN/m2)", 150.0)
        
        st.markdown("**Harga Satuan (RAB)**")
        p_semen = st.number_input("Semen (Rp/kg)", 1500)
        p_pasir = st.number_input("Pasir (Rp/m3)", 250000)
        p_split = st.number_input("Split (Rp/m3)", 300000)
        p_besi = st.number_input("Besi (Rp/kg)", 14000)
        p_kayu = st.number_input("Kayu (Rp/m3)", 2500000)
        p_batu = st.number_input("Batu Kali (Rp/m3)", 280000)
        p_beton_ready = st.number_input("Readymix K300", 1100000)
        p_bata = st.number_input("Bata Merah (bh)", 800)
        p_cat = st.number_input("Cat (kg)", 25000)
        p_pipa = st.number_input("Pipa 3/4 (m)", 15000)
        u_tukang = st.number_input("Tukang (OH)", 135000)
        u_pekerja = st.number_input("Pekerja (OH)", 110000)

# ==========================================
# 6. LOGIKA MODE: KALKULATOR TEKNIK
# ==========================================
if app_mode == "🧮 Kalkulator Teknik (Tools)":
    
    # Init Engines
    calc_sni = sni.SNI_Concrete_2847(fc_in, fy_in)
    calc_biaya = ahsp.AHSP_Engine()
    calc_geo = geo.Geotech_Engine(gamma_tanah, phi_tanah, c_tanah)
    calc_fdn = fdn.Foundation_Engine(sigma_tanah)
    engine_export = exp.Export_Engine()

    st.markdown(f'<div class="main-header">🛠️ Engineering Workspace: {nama_proyek}</div>', unsafe_allow_html=True)

    tabs = st.tabs([
        "🏠 Dash", "📂 BIM Import", "✏️ Model & Draw", "🏗️ Beton", 
        "🔩 Baja/Atap", "🌋 Gempa", "⛰️ Geoteknik", "💰 RAB Final"
    ])

    # [TAB 1: DASHBOARD]
    with tabs[0]:
        c1, c2, c3 = st.columns(3)
        c1.metric("Mutu Beton", f"{fc_in} MPa")
        c2.metric("Mutu Baja", f"{fy_in} MPa")
        c3.metric("Tanah (Phi)", f"{phi_tanah}°")
        st.info("Selamat Datang di IndoBIM Ultimate. Gunakan tab di atas untuk perhitungan teknik.")

    # [TAB 2: BIM IMPORT - DENGAN FITUR VISUALISASI 3D]
    with tabs[1]:
        st.markdown("### 📂 Upload Model 3D (.IFC)")
        uploaded_ifc = st.file_uploader("Upload File IFC", type=["ifc"])
        
        if uploaded_ifc:
            try:
                with st.spinner("🚀 Sedang membedah struktur 3D..."):
                    eng_ifc = bim.IFC_Parser_Engine(uploaded_ifc)
                    df_s = eng_ifc.parse_structure()
                    q_a = eng_ifc.parse_architectural_quantities()
                    q_m = eng_ifc.parse_mep_quantities()
                    
                    st.success(f"✅ Berhasil membaca {len(df_s)} elemen struktur!")
                    
                    # [TAMPILAN METRIK DATA]
                    st.markdown("#### 📊 Ringkasan Arsitektur")
                    col1, col2, col3 = st.columns(3)
                    
                    # Gunakan .get() untuk keamanan jika data kosong
                    luas_dinding = q_a.get('Luas Dinding (m2)', 0)
                    jml_pintu = q_a.get('Jumlah Pintu (Unit)', 0)
                    jml_jendela = q_a.get('Jumlah Jendela (Unit)', 0)
                    
                    col1.metric("🧱 Luas Dinding", f"{luas_dinding} m²")
                    col2.metric("🚪 Jumlah Pintu", f"{jml_pintu} Unit")
                    col3.metric("🪟 Jumlah Jendela", f"{jml_jendela} Unit")
                    
                    with st.expander("Lihat Data MEP (Pipa & Duct)"):
                        st.json(q_m)
                        
                    with st.expander("Lihat Tabel Elemen Struktur"):
                        st.dataframe(df_s.head(10))

                    # Tombol Simpan
                    if st.button("💾 Simpan Data ke Memori AI"):
                        st.session_state['arsitek_mep'] = {**q_a, **q_m}
                        st.session_state['bim_loads'] = eng_ifc.calculate_architectural_loads()['Total Load Tambahan (kN)']
                        st.toast("Data BIM tersimpan! Silakan chat dengan Konsultan AI.", icon="🤖")
                        st.balloons()
                    
                    # === [VISUALISASI 3D UNTUK DEBUG KOORDINAT] ===
                    st.divider()
                    st.markdown("#### 🕵️ Verifikasi Visual (3D Plot)")
                    st.caption("Gunakan ini untuk memastikan koordinat tidak 0,0,0 (menumpuk).")
                    
                    if st.checkbox("Tampilkan Preview Struktur (Scatter Plot)"):
                        if not df_s.empty:
                            # Cek apakah semua koordinat 0?
                            if df_s['X'].sum() == 0 and df_s['Y'].sum() == 0:
                                st.warning("⚠️ PERINGATAN: Semua koordinat terdeteksi (0,0,0). Isu pada 'Base Point' file IFC.")
                            
                            fig = plt.figure(figsize=(8, 6))
                            ax = fig.add_subplot(111, projection='3d')
                            
                            # Warna berdasarkan tipe
                            colors = {'Column': 'red', 'Beam': 'blue', 'Wall': 'gray', 'Slab': 'green', 'Footing': 'black'}
                            
                            count_obj = 0
                            for idx, row in df_s.iterrows():
                                c = colors.get(row['Type'], 'cyan')
                                ax.scatter(row['X'], row['Y'], row['Z'], c=c, marker='o', s=20)
                                count_obj += 1
                                if count_obj > 500: break # Limit render
                            
                            ax.set_xlabel('X (m)')
                            ax.set_ylabel('Y (m)')
                            ax.set_zlabel('Z (m)')
                            ax.set_title(f"Preview Posisi {count_obj} Elemen Pertama")
                            st.pyplot(fig)
                        else:
                            st.write("Belum ada data struktur.")

            except Exception as e: 
                st.error(f"Terjadi kesalahan saat parsing IFC: {e}")

    # [TAB 3: MODELING & CANVAS]
    with tabs[2]:
        sub1, sub2 = st.tabs(["A. Input Grid", "B. Gambar Canvas"])
        with sub1:
            c1, c2 = st.columns([1,2])
            with c1:
                L = st.number_input("Panjang (m)", 2.0, 20.0, st.session_state['geo']['L'])
                b = st.number_input("Lebar (mm)", 150, 1000, st.session_state['geo']['b'])
                h = st.number_input("Tinggi (mm)", 200, 2000, st.session_state['geo']['h'])
                st.session_state['geo'] = {'L': L, 'b': b, 'h': h}
            with c2:
                fig, ax = plt.subplots(figsize=(6,2))
                ax.add_patch(patches.Rectangle((0,0), L, h/1000, facecolor='#2E86C1'))
                ax.set_xlim(-0.5, L+0.5); ax.set_ylim(-0.5, 2); ax.set_aspect('equal')
                st.pyplot(fig)
        
        with sub2:
            st.info("Gambar denah ruangan di bawah:")
            scale = st.slider("Skala Px/m", 10, 50, 20)
            canvas = st_canvas(fill_color="rgba(0, 150, 255, 0.3)", stroke_width=2, height=300, width=600, drawing_mode="rect", key="cvs")
            
            if canvas.json_data:
                rooms = []
                for obj in canvas.json_data["objects"]:
                    w, h_ = obj["width"]/scale, obj["height"]/scale
                    rooms.append({"Keliling": 2*(w+h_), "Luas": w*h_})
                
                if rooms:
                    df_r = pd.DataFrame(rooms)
                    v_dinding = df_r["Keliling"].sum() * 3.5
                    v_beton = (df_r["Keliling"].sum() * 0.15 * 0.35)
                    st.success(f"Est. Dinding: {v_dinding:.1f} m2 | Beton: {v_beton:.1f} m3")
                    if st.button("Pakai Data Gambar"):
                        st.session_state['drawing'] = {'vol_dinding': v_dinding, 'vol_beton': v_beton}

    # [TAB 4: BETON]
    with tabs[3]:
        c1, c2 = st.columns(2)
        with c1:
            q_dl = st.number_input("DL (kN/m)", 0.0, 50.0, 15.0)
            if st.session_state.get('bim_loads'): st.caption(f"+ BIM: {st.session_state['bim_loads']} kN")
            q_ll = st.number_input("LL (kN/m)", 0.0, 50.0, 5.0)
        with c2:
            q_u = sni.SNI_Load_1727.komb_pembebanan(q_dl, q_ll)
            Mu = (1/8) * q_u * (st.session_state['geo']['L']**2)
            As_req = calc_sni.kebutuhan_tulangan(Mu, st.session_state['geo']['b'], st.session_state['geo']['h'], 40)
            dia = st.selectbox("D Tulangan", [13,16,19,22])
            n_bars = np.ceil(As_req / (0.25 * 3.14 * dia**2))
            
            st.metric("Mu", f"{Mu:.1f} kNm")
            st.success(f"Desain: {int(n_bars)} D{dia}")
            
            vol_c = st.session_state['geo']['L'] * (st.session_state['geo']['b']/1000) * (st.session_state['geo']['h']/1000)
            st.session_state['structure'] = {'vol_beton': vol_c, 'berat_besi': vol_c*150}
            st.session_state['report_struk'] = {'Mu': Mu, 'Tulangan': f"{int(n_bars)} D{dia}", 'Dimensi': f"{st.session_state['geo']['b']}x{st.session_state['geo']['h']}"}
            
            p_beam = {'b': st.session_state['geo']['b'], 'h': st.session_state['geo']['h'], 'dia': dia, 'n': n_bars, 'pjg': st.session_state['geo']['L']}
            st.download_button("📥 DXF Balok", engine_export.create_dxf("BALOK", p_beam), "balok.dxf")

    # [TAB 5: BAJA]
    with tabs[4]:
        sub1, sub2 = st.tabs(["WF", "Atap"])
        with sub1:
            mu_b = st.number_input("Mu Baja (kNm)", 50.0)
            lb_b = st.number_input("Lb (m)", 3.0)
            db_wf = {"WF 200": {'Zx': 213}, "WF 250": {'Zx': 324}}
            sel_wf = st.selectbox("Profil", list(db_wf.keys()))
            eng_st = steel.SNI_Steel_1729(fy_in, 410)
            res_st = eng_st.cek_balok_lentur(mu_b, db_wf[sel_wf], lb_b)
            if res_st['Ratio'] <= 1.0: st.success("Aman")
            else: st.error("Gagal")
            st.session_state['report_baja'] = {'Profil': sel_wf, 'Ratio': res_st['Ratio'], 'Status': res_st['Status']}
        with sub2:
            la = st.number_input("Luas Atap (m2)", 100.0)
            calc_tr = steel.Baja_Ringan_Calc()
            st.write(calc_tr.hitung_kebutuhan_atap(la, "Metal"))

    # [TAB 6: GEMPA]
    with tabs[5]:
        ss = st.number_input("Ss", 0.8); s1 = st.number_input("S1", 0.4)
        sc = st.selectbox("Site Class", ["SE", "SD", "SC"])
        eng_q = quake.SNI_Gempa_1726(ss, s1, sc)
        v_base, sds, sd1 = eng_q.hitung_base_shear(2000, 8.0)
        st.metric("Base Shear V", f"{v_base:.1f} kN")
        st.session_state['report_gempa'] = {'V_gempa': v_base}

    # [TAB 7: GEOTEK]
    with tabs[6]:
        sub1, sub2 = st.tabs(["Pondasi", "Talud"])
        with sub1:
            pu = st.number_input("Pu (kN)", 150.0)
            bf = st.number_input("Lebar (m)", 1.0)
            res_f = calc_fdn.hitung_footplate(pu, bf, bf, 300)
            if "AMAN" in res_f['status']: st.success("Aman")
            st.session_state['pondasi'] = {'fp_beton': res_f['vol_beton'], 'fp_besi': res_f['berat_besi'], 'galian': res_f['vol_galian']}
        with sub2:
            ht = st.number_input("H Talud (m)", 3.0)
            res_t = calc_geo.hitung_talud_batu_kali(ht, 0.4, 1.5)
            st.write(f"SF Guling: {res_t['SF_Guling']:.2f}")
            st.session_state['geotech'] = {'vol_talud': res_t['Vol_Per_M'], 'vol_pile': 0}
            st.session_state['report_geo'] = {'Talud_SF': res_t['SF_Geser'], 'Pile_Qall': '-'}

    # [TAB 8: RAB (FINAL INTEGRATION)]
    with tabs[7]:
        st.subheader("💰 Rekapitulasi Anggaran Biaya (RAB)")
        
        # 1. Ambil Data dari Modul Lain
        d_str = st.session_state.get('structure', {})
        d_pon = st.session_state.get('pondasi', {})
        d_geo = st.session_state.get('geotech', {})
        d_bim = st.session_state.get('arsitek_mep', {}) # Data IFC
        d_draw = st.session_state.get('drawing', {})
        
        # 2. Volume Struktur (Manual / Drawing)
        vol_beton = d_str.get('vol_beton', 0) + d_pon.get('fp_beton', 0) + d_draw.get('vol_beton', 0)
        vol_besi = d_str.get('berat_besi', 0) + d_pon.get('fp_besi', 0)
        
        # 3. Volume Arsitektur (Prioritas BIM > Drawing)
        vol_dinding_bim = d_bim.get('Luas Dinding (m2)', 0)
        vol_dinding_draw = d_draw.get('vol_dinding', 0)
        vol_dinding_final = max(vol_dinding_bim, vol_dinding_draw)
        
        # Data MEP & Bukaan dari BIM
        jml_pintu = d_bim.get('Jumlah Pintu (Unit)', 0)
        jml_jendela = d_bim.get('Jumlah Jendela (Unit)', 0)
        pjg_pipa = d_bim.get('Panjang Pipa/Duct (m\')', 0)
        
        # 4. Harga Dasar (Ambil dari Sidebar)
        h_mat = {'semen': p_semen, 'pasir': p_pasir, 'split': p_split, 'besi': p_besi, 'kayu': p_kayu, 'batu kali': p_batu, 'beton k300': p_beton_ready, 'bata merah': p_bata, 'cat tembok': p_cat, 'pipa pvc': p_pipa}
        h_wage = {'pekerja': u_pekerja, 'tukang': u_tukang, 'mandor': u_pekerja*1.2}
        
        # 5. Hitung HSP (Analisa Harga Satuan Pekerjaan)
        hsp_beton = calc_biaya.hitung_hsp('beton_k250', h_mat, h_wage)
        hsp_besi = calc_biaya.hitung_hsp('pembesian_polos', h_mat, h_wage) / 10
        hsp_bata = calc_biaya.hitung_hsp('pasangan_bata_merah', h_mat, h_wage)
        hsp_plester = calc_biaya.hitung_hsp('plesteran', h_mat, h_wage)
        hsp_acian = calc_biaya.hitung_hsp('acian', h_mat, h_wage)
        hsp_cat = calc_biaya.hitung_hsp('cat_tembok', h_mat, h_wage)
        hsp_talud = calc_biaya.hitung_hsp('pasangan_batu_kali', h_mat, h_wage)
        hsp_pile = calc_biaya.hitung_hsp('bore_pile_k300', h_mat, h_wage)
        hsp_kusen = calc_biaya.hitung_hsp('pasang_kus_pintu', h_mat, h_wage)
        hsp_pipa = calc_biaya.hitung_hsp('pasang_pipa_pvc', h_mat, h_wage)
        
        # 6. Susun Tabel RAB
        rab_data = []
        
        # A. PEKERJAAN STRUKTUR
        if vol_beton > 0: rab_data.append({"Item": "Beton Struktur (K250)", "Vol": vol_beton, "Sat": "m3", "Hrg": hsp_beton, "Tot": vol_beton*hsp_beton})
        if vol_besi > 0: rab_data.append({"Item": "Pembesian Polos", "Vol": vol_besi, "Sat": "kg", "Hrg": hsp_besi, "Tot": vol_besi*hsp_besi})
        if d_geo.get('vol_talud', 0) > 0: rab_data.append({"Item": "Pondasi Talud Batu Kali", "Vol": d_geo['vol_talud'], "Sat": "m3", "Hrg": hsp_talud, "Tot": d_geo['vol_talud']*hsp_talud})
        if d_geo.get('vol_pile', 0) > 0: rab_data.append({"Item": "Pondasi Bore Pile", "Vol": d_geo['vol_pile'], "Sat": "m3", "Hrg": hsp_pile, "Tot": d_geo['vol_pile']*hsp_pile})
        
        # B. PEKERJAAN ARSITEKTUR (Otomatis dari BIM)
        if vol_dinding_final > 0:
            rab_data.append({"Item": "Pas. Dinding Bata Merah", "Vol": vol_dinding_final, "Sat": "m2", "Hrg": hsp_bata, "Tot": vol_dinding_final*hsp_bata})
            rab_data.append({"Item": "Plesteran Dinding", "Vol": vol_dinding_final*2, "Sat": "m2", "Hrg": hsp_plester, "Tot": vol_dinding_final*2*hsp_plester})
            rab_data.append({"Item": "Acian Dinding", "Vol": vol_dinding_final*2, "Sat": "m2", "Hrg": hsp_acian, "Tot": vol_dinding_final*2*hsp_acian})
            rab_data.append({"Item": "Pengecatan Dinding", "Vol": vol_dinding_final*2, "Sat": "m2", "Hrg": hsp_cat, "Tot": vol_dinding_final*2*hsp_cat})
        
        if jml_pintu > 0:
            rab_data.append({"Item": "Pasang Kusen Pintu/Jendela", "Vol": jml_pintu+jml_jendela, "Sat": "Unit", "Hrg": hsp_kusen, "Tot": (jml_pintu+jml_jendela)*hsp_kusen})
            
        # C. PEKERJAAN MEP (Otomatis dari BIM)
        if pjg_pipa > 0:
            rab_data.append({"Item": "Instalasi Pipa PVC", "Vol": pjg_pipa, "Sat": "m'", "Hrg": hsp_pipa, "Tot": pjg_pipa*hsp_pipa})
            
        # 7. TAMPILKAN HASIL
        if rab_data:
            df_rab = pd.DataFrame(rab_data)
            st.dataframe(df_rab.style.format({"Vol": "{:.2f}", "Hrg": "{:,.0f}", "Tot": "{:,.0f}"}), use_container_width=True)
            
            total_biaya = df_rab['Tot'].sum()
            st.success(f"### 🏷️ TOTAL ESTIMASI BIAYA: Rp {total_biaya:,.0f}")
            
            if vol_dinding_bim > 0:
                st.info(f"💡 Info: Data Arsitektur (Dinding {vol_dinding_bim} m2) & MEP diambil otomatis dari model BIM.")
            
            # Download Excel
            s_data = {'fc': fc_in, 'fy': fy_in, 'b': st.session_state['geo']['b'], 'h': st.session_state['geo']['h'], 'sigma': sigma_tanah}
            st.download_button("📊 Download Excel RAB", engine_export.create_excel_report(df_rab, s_data), "RAB_Project.xlsx")
        else:
            st.warning("Belum ada volume pekerjaan yang terhitung. Silakan input data di tab Struktur atau Upload BIM.")

# ==========================================
# 6. LOGIKA MODE: KONSULTAN AI (ENGINEX)
# ==========================================
elif app_mode == "🤖 Konsultan AI (Chat)":
    
    st.markdown(f'<div class="main-header">🤖 AI Consultant: {nama_proyek}</div>', unsafe_allow_html=True)
    
    # --- PILIH AHLI (SIDEBAR BAWAH) ---
    with st.sidebar:
        st.markdown("---")
        st.markdown("### 👷 Tim Ahli")
        use_auto = st.checkbox("🤖 Auto-Pilot", value=True)
        manual_sel = st.selectbox("Pilih Manual:", list(gems_persona.keys()), disabled=use_auto)
        if not use_auto: st.session_state.current_expert_active = manual_sel
        
        # File Upload
        st.markdown("### 📂 Upload Data")
        uploaded_files = st.file_uploader("File Pendukung:", accept_multiple_files=True)
        if st.button("🧹 Reset Chat"):
            db.clear_chat(nama_proyek, st.session_state.current_expert_active)
            st.rerun()

    # --- CHAT AREA ---
    current_expert = st.session_state.current_expert_active
    st.caption(f"Status: **Connected** | Expert: **{current_expert}** | Brain: **{selected_model_name}**")
    
    # Render History
    history = db.get_chat_history(nama_proyek, current_expert)
    for chat in history:
        with st.chat_message(chat['role']): st.markdown(chat['content'])
        
    if prompt := st.chat_input("Tanya sesuatu..."):
        # Auto-Pilot Logic
        target_expert = current_expert
        if use_auto:
            target_expert = get_auto_pilot_decision(prompt, selected_model_name)
            st.session_state.current_expert_active = target_expert
            st.toast(f"Dialihkan ke: {target_expert}", icon="🔀")
            
        # Save User Msg
        db.simpan_chat(nama_proyek, target_expert, "user", prompt)
        with st.chat_message("user"): st.markdown(prompt)
        
        # Prepare Context
        context_files = ""
        if uploaded_files:
            # Simplifikasi: Hanya memberi tahu AI ada file
            context_files = "\n[Ada File Terlampir]"
            
        final_context = get_project_summary_context() + context_files + "\n\n" + prompt
        
        # Generate Response
        with st.chat_message("assistant"):
            with st.spinner("Berpikir..."):
                try:
                    model = genai.GenerativeModel(selected_model_name, system_instruction=gems_persona[target_expert])
                    
                    # Convert history for API
                    api_hist = [{"role": "user" if h['role']=="user" else "model", "parts": [h['content']]} for h in history]
                    
                    chat = model.start_chat(history=api_hist)
                    response = chat.send_message(final_context)
                    ans = response.text
                    
                    st.markdown(ans)
                    db.simpan_chat(nama_proyek, target_expert, "assistant", ans)
                    
                    # Cek Plotting Code
                    if "```python" in ans and "plt." in ans:
                        code = re.search(r"```python(.*?)```", ans, re.DOTALL).group(1)
                        st.caption("📈 Rendering Grafik...")
                        execute_generated_code(code)
                        
                    # Download Docs
                    docx_bio = create_docx_from_text(ans)
                    if docx_bio: st.download_button("📄 Simpan Word", docx_bio, "Jawaban.docx")
                    
                except Exception as e:
                    st.error(f"Error AI: {e}")
